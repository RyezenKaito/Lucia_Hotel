from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_class_table(doc, index, class_name, attributes, methods):
    doc.add_heading(f'Table {index}. Thực thể {class_name}', level=3)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Công việc / Thuộc tính'
    hdr_cells[2].text = 'Kiểu dữ liệu'
    hdr_cells[3].text = 'Ràng buộc'
    hdr_cells[4].text = 'Ghi chú'

    # Attributes
    row_cells = table.add_row().cells
    row_cells[0].text = '1'
    row_cells[1].text = 'Khai báo thuộc tính:'
    
    for i, attr in enumerate(attributes):
        row_cells = table.add_row().cells
        row_cells[0].text = f'1.{i+1}'
        row_cells[1].text = attr[0]
        row_cells[2].text = attr[1]
        row_cells[3].text = attr[2]
        row_cells[4].text = attr[3]

    # Methods
    if methods:
        row_cells = table.add_row().cells
        row_cells[0].text = '2'
        row_cells[1].text = 'Viết các phương thức'
        for i, meth in enumerate(methods):
            row_cells = table.add_row().cells
            row_cells[0].text = f'2.{i+1}'
            row_cells[1].text = meth[0]
            row_cells[2].text = meth[1]
            row_cells[3].text = meth[2]
            row_cells[4].text = meth[3]

def add_db_table(doc, index, table_name, columns):
    doc.add_heading(f'3.2.{index} Bảng {table_name}', level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên thuộc tính'
    hdr_cells[1].text = 'Kiểu dữ liệu'
    hdr_cells[2].text = 'Khóa (PK/FK)'
    hdr_cells[3].text = 'Ghi chú / Ràng buộc'

    for col in columns:
        row_cells = table.add_row().cells
        row_cells[0].text = col[0]
        row_cells[1].text = col[1]
        row_cells[2].text = col[2]
        row_cells[3].text = col[3]

doc = Document()
doc.add_heading('CHƯƠNG 3: THIẾT KẾ VÀ HIỆN THỰC', level=1)

doc.add_heading('3.1.2 Đặc tả lớp', level=2)

add_class_table(doc, 1, 'Nhân viên (NV)', 
    [('maNV', 'String', "Bắt đầu bằng 'LUCIA' và đi kèm số hoặc là 'ADMIN'", '{id}'), 
     ('hoTen', 'String', 'Không được rỗng', ''), 
     ('soCCCD', 'String', 'Gồm 9 hoặc 12 chữ số, UNIQUE', ''), 
     ('soDT', 'String', 'Đủ 10 chữ số và bắt đầu từ số 0, UNIQUE', ''), 
     ('diaChi', 'String', '', ''), 
     ('ngayVaoLam', 'Date', 'Nhỏ hơn hoặc bằng ngày hiện tại', ''), 
     ('ngaySinh', 'Date', 'Phải đủ 18 tuổi trở lên', ''), 
     ('trinhDo', 'trinhDo', "IN ('THCS', 'THPT', 'CAODANG', 'DAIHOC')", 'Enum'), 
     ('mk', 'String', 'Không được rỗng', ''), 
     ('role', 'role', 'Không được để trống', 'Enum'), 
     ('trangThai', 'trangThai', 'Không được để trống', 'Enum'), 
     ('maQL', 'NV', "Có thể null, 'LUCIA' kèm số hoặc 'ADMIN'", 'Foreign Key')], 
    [('setMaNV(String)', 'void', "Ném IllegalArgumentException nếu sai format", ""),
     ('setNgaySinh(LocalDate)', 'void', "Khoảng cách năm so với hiện tại >= 18", ""),
     ('setSoDT(String)', 'void', "REGEX 0\\d{9}", ""),
     ('setCccd(String)', 'void', "REGEX \\d{9}(\\d{3})?", ""),
     ('setNgayVaoLamDate(LocalDate)', 'void', "Không được sau ngày hiện tại", "")])

add_class_table(doc, 2, 'Khách hàng (KH)', 
    [('maKH', 'String', 'NOT NULL', '{id}'), 
     ('tenKH', 'String', 'NOT NULL', ''), 
     ('soDT', 'String', 'UNIQUE', ''), 
     ('ngaySinh', 'Date', '', ''), 
     ('soCCCD', 'String', 'UNIQUE', '')], 
    [])

add_class_table(doc, 3, 'Phòng (Phong)', 
    [('maPhong', 'String', 'NOT NULL', '{id}'), 
     ('tenPhong', 'String', 'NOT NULL', ''), 
     ('loaiPhong', 'LoaiPhong', 'NOT NULL', 'Foreign Key'), 
     ('tinhTrang', 'tinhTrang', "IN ('BAN', 'CONTRONG', 'DANGSUDUNG')", 'Enum'), 
     ('soPhong', 'Int', 'NOT NULL', ''), 
     ('soTang', 'Int', 'NOT NULL', '')], 
    [('getDanhSach(Phong)', 'void', '', ''), 
     ('checkStatus(tinhTrang)', 'void', '', ''), 
     ('updateStatus(tinhTrang)', 'void', '', '')])

add_class_table(doc, 4, 'Đặt Phòng (DatPhong)', 
    [('maDat', 'String', 'NOT NULL', '{id}'), 
     ('ngayDat', 'Date', 'NOT NULL', ''), 
     ('maKH', 'KH', 'NOT NULL', 'Foreign Key'), 
     ('ngayCheckIn', 'Date', 'ngayCheckIn >= ngayDat', ''), 
     ('ngayCheckOut', 'Date', 'ngayCheckOut >= ngayCheckIn', ''), 
     ('trangThai', 'String', 'NOT NULL', '')], 
    [('setCheck_inTime(Date, Time)', 'void', '', ''), 
     ('setCheck_outTime(Date, Time)', 'void', '', ''), 
     ('checkBooking(maKH)', 'void', '', '')])

add_class_table(doc, 5, 'Hóa Đơn (HoaDon)', 
    [('maHD', 'String', 'NOT NULL', '{id}'), 
     ('maDat', 'DatPhong', 'NOT NULL', 'Foreign Key'), 
     ('maNV', 'NV', 'NOT NULL', 'Foreign Key'), 
     ('ngayTaoHD', 'Date', 'NOT NULL', ''), 
     ('tienPhong', 'Double', '>= 0', ''), 
     ('VAT', 'Double', '>= 0', ''), 
     ('tongCoc', 'Double', '>= 0', ''), 
     ('tienDV', 'Double', '>= 0', ''), 
     ('tongTien', 'Double', '>= 0', ''), 
     ('doanhThu', 'Double', '>= 0', ''), 
     ('loaiHD', 'String', '', ''), 
     ('trangThaiThanhToan', 'String', '', ''), 
     ('ngayThanhToan', 'Date', '', ''), 
     ('ghiChuThanhToan', 'String', '', '')], 
    [('updateStatus()', 'void', '', ''), 
     ('saveTTHD(maHD, ngayTaoHD)', 'void', '', ''), 
     ('tongTien()', 'void', '', ''), 
     ('operation()', 'void', '', '')])

add_class_table(doc, 6, 'Chi Tiết Đặt Phòng (ChiTietDatPhong)', 
    [('maCTDP', 'String', 'NOT NULL', '{id}'), 
     ('maPhong', 'Phong', 'NOT NULL', 'Foreign Key'), 
     ('maDat', 'DatPhong', 'NOT NULL', 'Foreign Key'), 
     ('giaCoc', 'Double', '>= 0', ''), 
     ('soNguoi', 'Int', '> 0', ''), 
     ('ghiChu', 'String', '', '')], 
    [])

add_class_table(doc, 7, 'Chi Tiết Hóa Đơn (ChiTietHoaDon)', 
    [('maCTHD', 'String', 'NOT NULL', '{id}'), 
     ('maHD', 'HoaDon', 'NOT NULL', 'Foreign Key'), 
     ('maCTDP', 'ChiTietDatPhong', 'NOT NULL', 'Foreign Key'), 
     ('thoiGianLuuTru', 'Double', '>= 0', ''), 
     ('soLuongPhong', 'Int', '> 0', ''), 
     ('thanhTien', 'Double', '>= 0', '')], 
    [('getTTHD(maHD)', 'void', '', ''), 
     ('tongTien()', 'void', '', '')])

add_class_table(doc, 8, 'Dịch Vụ (DV)', 
    [('maDV', 'String', 'NOT NULL', '{id}'), 
     ('tenDV', 'String', 'NOT NULL', ''), 
     ('gia', 'Double', '>= 0', ''), 
     ('loaiDV', 'loaiDV', 'NOT NULL', 'Enum'), 
     ('mieuTa', 'String', '', ''), 
     ('donVi', 'String', 'NOT NULL', ''), 
     ('trangThai', 'Boolean', 'NOT NULL', '')], 
    [])

add_class_table(doc, 9, 'Bảng Giá Dịch Vụ Header (BangGiaDV_Header)', 
    [('maBangGia', 'String', 'NOT NULL', '{id}'), 
     ('tenBangGia', 'String', 'NOT NULL', ''), 
     ('ngayApDung', 'Date', 'NOT NULL', ''), 
     ('ngayHetHieuLuc', 'Date', 'ngayHetHieuLuc > ngayApDung', ''), 
     ('trangThai', 'Boolean', 'NOT NULL', '')], 
    [])

add_class_table(doc, 10, 'Bảng Giá Dịch Vụ Detail (BangGiaDV_Detail)', 
    [('maBangGia', 'BangGiaDV_Header', 'NOT NULL', '{id}, Foreign Key'), 
     ('maDV', 'DV', 'NOT NULL', 'Foreign Key'), 
     ('giaDV', 'Double', '>= 0', '')], 
    [])

add_class_table(doc, 11, 'Dịch Vụ Sử Dụng (DichVuSuDung)', 
    [('maDV', 'DV', 'NOT NULL', '{id}, Foreign Key'), 
     ('maCTDP', 'ChiTietDatPhong', 'NOT NULL', 'Foreign Key'), 
     ('ngaySuDung', 'Date', 'NOT NULL', ''), 
     ('soLuong', 'Int', '> 0', ''), 
     ('giaDV', 'Double', '>= 0', ''), 
     ('trangThai', 'Boolean', 'NOT NULL', '')], 
    [('thanhTien(giaDV)', 'Double', '', '')])

add_class_table(doc, 12, 'Loại Phòng (LoaiPhong)', 
    [('maLoai', 'String', 'NOT NULL', '{id}'), 
     ('donGia', 'Double', '> 0', ''), 
     ('sucChua', 'Int', '> 0', '')], 
    [])


doc.add_heading('3.2.3 Các ràng buộc toàn vẹn trong CSDL', level=2)

add_db_table(doc, 1, 'Khách hàng (KH)', [
    ('maKH', 'VARCHAR', 'PK', 'NOT NULL'),
    ('tenKH', 'NVARCHAR', '', 'NOT NULL'),
    ('soDT', 'VARCHAR', '', 'UNIQUE'),
    ('ngaySinh', 'DATE', '', ''),
    ('soCCCD', 'VARCHAR', '', 'UNIQUE')
])

add_db_table(doc, 2, 'Nhân viên (NV)', [
    ('maNV', 'VARCHAR', 'PK', 'NOT NULL'),
    ('hoTen', 'NVARCHAR', '', 'NOT NULL'),
    ('soDT', 'VARCHAR', '', 'UNIQUE'),
    ('soCCCD', 'VARCHAR', '', 'UNIQUE'),
    ('diaChi', 'NVARCHAR', '', ''),
    ('ngaySinh', 'DATE', '', 'ngaySinh < ngayVaoLam - 18 năm'),
    ('ngayVaoLam', 'DATE', '', ''),
    ('trinhDo', 'NVARCHAR', '', "IN ('THCS', 'THPT', 'CAODANG', 'DAIHOC')"),
    ('mk', 'VARCHAR', '', 'NOT NULL'),
    ('role', 'NVARCHAR', '', "IN ('NV', 'QL')"),
    ('trangThai', 'NVARCHAR', '', "IN ('CONLAM', 'DANGHI')"),
    ('maQL', 'VARCHAR', 'FK', 'REFERENCES NV(maNV)')
])

add_db_table(doc, 3, 'Loại phòng (LoaiPhong)', [
    ('maLoaiPhong', 'VARCHAR', 'PK', 'NOT NULL'),
    ('gia', 'DECIMAL', '', 'NOT NULL, > 0'),
    ('sucChua', 'INT', '', 'NOT NULL, > 0')
])

add_db_table(doc, 4, 'Phòng (Phong)', [
    ('maPhong', 'VARCHAR', 'PK', 'NOT NULL'),
    ('tenPhong', 'NVARCHAR', '', 'NOT NULL'),
    ('loaiPhong', 'VARCHAR', 'FK', 'REFERENCES LoaiPhong(maLoaiPhong)'),
    ('tinhTrang', 'NVARCHAR', '', "IN ('BAN', 'CONTRONG', 'DANGSUDUNG')"),
    ('soPhong', 'INT', '', 'NOT NULL'),
    ('soTang', 'INT', '', 'NOT NULL')
])

add_db_table(doc, 5, 'Đặt Phòng (DatPhong)', [
    ('maDat', 'VARCHAR', 'PK', 'NOT NULL'),
    ('ngayDat', 'DATETIME', '', 'NOT NULL'),
    ('maKH', 'VARCHAR', 'FK', 'REFERENCES KH(maKH)'),
    ('ngayCheckIn', 'DATETIME', '', 'NOT NULL, >= ngayDat'),
    ('ngayCheckOut', 'DATETIME', '', 'NOT NULL, >= ngayCheckIn'),
    ('trangThai', 'NVARCHAR', '', 'NOT NULL')
])

add_db_table(doc, 6, 'Chi Tiết Đặt Phòng (ChiTietDatPhong)', [
    ('maCTDP', 'VARCHAR', 'PK', 'NOT NULL'),
    ('maPhong', 'VARCHAR', 'FK', 'REFERENCES Phong(maPhong)'),
    ('maDat', 'VARCHAR', 'FK', 'REFERENCES DatPhong(maDat)'),
    ('giaCoc', 'DECIMAL', '', '>= 0'),
    ('soNguoi', 'INT', '', '> 0'),
    ('ghiChu', 'NVARCHAR', '', '')
])

add_db_table(doc, 7, 'Hóa Đơn (HoaDon)', [
    ('maHD', 'VARCHAR', 'PK', 'NOT NULL'),
    ('maDat', 'VARCHAR', 'FK', 'REFERENCES DatPhong(maDat)'),
    ('maNV', 'VARCHAR', 'FK', 'REFERENCES NV(maNV)'),
    ('ngayTaoHD', 'DATETIME', '', 'NOT NULL'),
    ('tienPhong', 'DECIMAL', '', '>= 0'),
    ('tienDV', 'DECIMAL', '', '>= 0'),
    ('tienCoc', 'DECIMAL', '', '>= 0'),
    ('thueVAT', 'DECIMAL', '', '>= 0'),
    ('tongTien', 'DECIMAL', '', '>= 0'),
    ('doanhThu', 'DECIMAL', '', '>= 0'),
    ('loaiHD', 'NVARCHAR', '', ''),
    ('trangThaiThanhToan', 'NVARCHAR', '', ''),
    ('phuongThucThanhToan', 'NVARCHAR', '', ''),
    ('ngayThanhToan', 'DATETIME', '', ''),
    ('ghiChuThanhToan', 'NVARCHAR', '', '')
])

add_db_table(doc, 8, 'Chi Tiết Hóa Đơn (ChiTietHoaDon)', [
    ('maCTHD', 'VARCHAR', 'PK', 'NOT NULL'),
    ('maHD', 'VARCHAR', 'FK', 'REFERENCES HoaDon(maHD)'),
    ('maCTDP', 'VARCHAR', 'FK', 'REFERENCES ChiTietDatPhong(maCTDP)'),
    ('thoiGianLuuTru', 'DECIMAL', '', '>= 0'),
    ('thanhTien', 'DECIMAL', '', '>= 0')
])

add_db_table(doc, 9, 'Dịch vụ (DV)', [
    ('maDV', 'VARCHAR', 'PK', 'NOT NULL'),
    ('tenDV', 'NVARCHAR', '', 'NOT NULL'),
    ('gia', 'DECIMAL', '', '>= 0'),
    ('loaiDV', 'NVARCHAR', '', ''),
    ('mieuTa', 'NVARCHAR', '', ''),
    ('donVi', 'NVARCHAR', '', 'NOT NULL'),
    ('trangThai', 'BIT', '', 'NOT NULL')
])

add_db_table(doc, 10, 'Bảng Giá Dịch Vụ - Header (BangGiaDV_Header)', [
    ('maBangGia', 'VARCHAR', 'PK', 'NOT NULL'),
    ('tenBangGia', 'NVARCHAR', '', ''),
    ('ngayApDung', 'DATE', '', 'NOT NULL'),
    ('ngayHetHieuLuc', 'DATE', '', 'ngayHetHieuLuc > ngayApDung'),
    ('trangThai', 'BIT', '', 'NOT NULL')
])

add_db_table(doc, 11, 'Bảng Giá Dịch Vụ - Detail (BangGiaDV_Detail)', [
    ('maBangGia', 'VARCHAR', 'PK, FK', 'REFERENCES BangGiaDV_Header(maBangGia)'),
    ('maDV', 'VARCHAR', 'PK, FK', 'REFERENCES DV(maDV)'),
    ('giaDV', 'DECIMAL', '', '>= 0')
])

add_db_table(doc, 12, 'Dịch vụ Sử Dụng (DichVuSuDung)', [
    ('maDV', 'VARCHAR', 'PK, FK', 'REFERENCES DV(maDV)'),
    ('maCTDP', 'VARCHAR', 'PK, FK', 'REFERENCES ChiTietDatPhong(maCTDP)'),
    ('ngaySuDung', 'DATE', '', 'NOT NULL'),
    ('soLuong', 'INT', '', '> 0'),
    ('giaDV', 'DECIMAL', '', '>= 0'),
    ('trangThai', 'BIT', '', 'NOT NULL')
])

doc.save('DacTa_RangBuoc_v2.docx')
print("Document generated successfully.")
