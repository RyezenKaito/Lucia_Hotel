from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_class_table(doc, index, class_name, attributes, methods):
    doc.add_heading(f'Table {index}. Thực thể {class_name}', level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'STT'
    hdr_cells[1].text = 'Công việc / Thuộc tính'
    hdr_cells[2].text = 'Kiểu dữ liệu'
    hdr_cells[3].text = 'Ghi chú'

    # Attributes
    row_cells = table.add_row().cells
    row_cells[0].text = '1'
    row_cells[1].text = 'Khai báo thuộc tính:'
    
    for i, attr in enumerate(attributes):
        row_cells = table.add_row().cells
        row_cells[0].text = f'1.{i+1}'
        row_cells[1].text = attr[0]
        row_cells[2].text = attr[1]
        row_cells[3].text = attr[2]

    # Methods
    if methods:
        row_cells = table.add_row().cells
        row_cells[0].text = '2'
        row_cells[1].text = 'Viết các phương thức'
        for i, meth in enumerate(methods):
            row_cells = table.add_row().cells
            row_cells[0].text = f'2.{i+1}'
            row_cells[1].text = meth[0]
            row_cells[2].text = meth[1]
            row_cells[3].text = meth[2]

def add_db_table(doc, index, table_name, columns):
    doc.add_heading(f'3.2.{index} Bảng {table_name}', level=3)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên thuộc tính'
    hdr_cells[1].text = 'Kiểu dữ liệu'
    hdr_cells[2].text = 'Khóa (PK/FK)'
    hdr_cells[3].text = 'Ghi chú / Ràng buộc'

    for col in columns:
        row_cells = table.add_row().cells
        row_cells[0].text = col[0]
        row_cells[1].text = col[1]
        row_cells[2].text = col[2]
        row_cells[3].text = col[3]

doc = Document()
doc.add_heading('CHƯƠNG 3: THIẾT KẾ VÀ HIỆN THỰC', level=1)

doc.add_heading('3.1.2 Đặc tả lớp', level=2)

add_class_table(doc, 1, 'Nhân viên (NV)', 
    [('maNV', 'String', '{id}'), ('hoTen', 'String', ''), ('soCCCD', 'String', ''), ('soDT', 'String', ''), 
     ('diaChi', 'String', ''), ('ngayVaoLam', 'Date', ''), ('ngaySinh', 'Date', ''), ('trinhDo', 'trinhDo', 'Enum'), 
     ('mk', 'String', ''), ('role', 'role', 'Enum'), ('trangThai', 'trangThai', 'Enum'), ('maQL', 'NV', '')], [])

add_class_table(doc, 2, 'Khách hàng (KH)', 
    [('maKH', 'String', '{id}'), ('tenKH', 'String', ''), ('soDT', 'String', ''), ('ngaySinh', 'Date', ''), ('soCCCD', 'String', '')], [])

add_class_table(doc, 3, 'Phòng (Phong)', 
    [('maPhong', 'String', '{id}'), ('tenPhong', 'String', ''), ('loaiPhong', 'LoaiPhong_Header', ''), 
     ('tinhTrang', 'tinhTrang', 'Enum'), ('soPhong', 'Int', ''), ('soTang', 'Int', ''), ('maLoai', 'LoaiPhong', '')], 
    [('getDanhSach(Phong)', 'void', ''), ('checkStatus(tinhTrang)', 'void', ''), ('updateStatus(tinhTrang)', 'void', '')])

add_class_table(doc, 4, 'Đặt Phòng (DatPhong)', 
    [('maDat', 'String', '{id}'), ('ngayDat', 'Date', ''), ('maKH', 'KH', ''), 
     ('ngayCheckIn', 'Date', ''), ('ngayCheckOut', 'Date', ''), ('trangThai', 'String', '')], 
    [('setCheck_inTime(Date, Time)', 'void', ''), ('setCheck_outTime(Date, Time)', 'void', ''), ('checkBooking(maKH)', 'void', '')])

add_class_table(doc, 5, 'Hóa Đơn (HoaDon)', 
    [('maHD', 'String', '{id}'), ('maDat', 'DatPhong', ''), ('maNV', 'NV', ''), ('ngayTaoHD', 'Date', ''), 
     ('tienPhong', 'Double', ''), ('VAT', 'Double', ''), ('tongCoc', 'Double', ''), ('tienDV', 'Double', ''), 
     ('tongTien', 'Double', ''), ('doanhThu', 'Double', ''), ('loaiHD', 'String', ''), 
     ('trangThaiThanhToan', 'String', ''), ('ngayThanhToan', 'Date', ''), ('ghiChuThanhToan', 'String', '')], 
    [('updateStatus()', 'void', ''), ('saveTTHD(maHD, ngayTaoHD)', 'void', ''), ('tongTien()', 'void', ''), ('operation()', 'void', '')])

add_class_table(doc, 6, 'Chi Tiết Đặt Phòng (ChiTietDatPhong)', 
    [('maCTDP', 'String', '{id}'), ('maPhong', 'Phong', ''), ('maDat', 'DatPhong', ''), 
     ('giaCoc', 'Double', ''), ('soNguoi', 'Int', ''), ('ghiChu', 'String', '')], [])

add_class_table(doc, 7, 'Chi Tiết Hóa Đơn (ChiTietHoaDon)', 
    [('maCTHD', 'String', '{id}'), ('maHD', 'HoaDon', ''), ('maCTDP', 'ChiTietDatPhong', ''), 
     ('thoiGianLuuTru', 'Double', ''), ('soLuongPhong', 'Int', ''), ('thanhTien', 'Double', '')], 
    [('getTTHD(maHD)', 'void', ''), ('tongTien()', 'void', '')])

add_class_table(doc, 8, 'Dịch Vụ (DV)', 
    [('maDV', 'String', '{id}'), ('tenDV', 'String', ''), ('gia', 'Double', ''), 
     ('loaiDV', 'loaiDV', 'Enum'), ('mieuTa', 'String', ''), ('donVi', 'String', ''), ('trangThai', 'String', '')], [])

add_class_table(doc, 9, 'Bảng Giá Dịch Vụ Header (BangGiaDV_Header)', 
    [('maBangGia', 'String', '{id}'), ('tenBangGia', 'String', ''), ('ngayApDung', 'Date', ''), 
     ('ngayHetHieuLuc', 'Date', ''), ('trangThai', 'Boolean', '')], [])

add_class_table(doc, 10, 'Bảng Giá Dịch Vụ Detail (BangGiaDV_Detail)', 
    [('maBangGia', 'BangGiaDV_Header', '{id}'), ('maDV', 'DV', ''), ('giaDV', 'Double', '')], [])

add_class_table(doc, 11, 'Dịch Vụ Sử Dụng (DichVuSuDung)', 
    [('maDV', 'DV', '{id}'), ('maCTHD', 'ChiTietHoaDon', ''), ('ngaySuDung', 'Date', ''), 
     ('soLuong', 'Int', ''), ('giaDV', 'Double', ''), ('trangThai', 'Boolean', '')], 
    [('thanhTien(giaDV)', 'Double', '')])

add_class_table(doc, 12, 'Loại Phòng (LoaiPhong)', 
    [('maLoai', 'String', '{id}'), ('donGia', 'Double', ''), ('sucChua', 'Int', '')], [])


doc.add_heading('3.2.3 Các ràng buộc toàn vẹn trong CSDL', level=2)

add_db_table(doc, 1, 'Khách hàng (KH)', [
    ('maKH', 'VARCHAR', 'PK', 'NOT NULL'),
    ('tenKH', 'NVARCHAR', '', 'NOT NULL'),
    ('soDT', 'VARCHAR', '', 'UNIQUE'),
    ('ngaySinh', 'DATE', '', ''),
    ('soCCCD', 'VARCHAR', '', 'UNIQUE')
])

add_db_table(doc, 2, 'Nhân viên (NV)', [
    ('maNV', 'VARCHAR', 'PK', 'NOT NULL'),
    ('hoTen', 'NVARCHAR', '', 'NOT NULL'),
    ('soDT', 'VARCHAR', '', 'UNIQUE'),
    ('soCCCD', 'VARCHAR', '', 'UNIQUE'),
    ('diaChi', 'NVARCHAR', '', ''),
    ('ngaySinh', 'DATE', '', ''),
    ('ngayVaoLam', 'DATE', '', ''),
    ('trinhDo', 'NVARCHAR', '', "IN ('THCS', 'THPT', 'CAODANG', 'DAIHOC')"),
    ('mk', 'VARCHAR', '', 'NOT NULL'),
    ('role', 'NVARCHAR', '', "IN ('NV', 'QL')"),
    ('trangThai', 'NVARCHAR', '', "IN ('CONLAM', 'DANGHI')"),
    ('maQL', 'VARCHAR', 'FK', 'REFERENCES NV(maNV)')
])

add_db_table(doc, 3, 'Loại phòng (LoaiPhong)', [
    ('maLoaiPhong', 'VARCHAR', 'PK', 'NOT NULL'),
    ('gia', 'DECIMAL', '', 'NOT NULL'),
    ('sucChua', 'INT', '', 'NOT NULL')
])

add_db_table(doc, 4, 'Phòng (Phong)', [
    ('maPhong', 'VARCHAR', 'PK', 'NOT NULL'),
    ('tenPhong', 'NVARCHAR', '', 'NOT NULL'),
    ('loaiPhong', 'VARCHAR', 'FK', 'REFERENCES LoaiPhong(maLoaiPhong)'),
    ('tinhTrang', 'NVARCHAR', '', "IN ('BAN', 'CONTRONG', 'DANGSUDUNG')"),
    ('soPhong', 'INT', '', 'NOT NULL'),
    ('soTang', 'INT', '', 'NOT NULL')
])

add_db_table(doc, 5, 'Đặt Phòng (DatPhong)', [
    ('maDat', 'VARCHAR', 'PK', 'NOT NULL'),
    ('ngayDat', 'DATETIME', '', 'NOT NULL'),
    ('maKH', 'VARCHAR', 'FK', 'REFERENCES KH(maKH)'),
    ('ngayCheckIn', 'DATETIME', '', 'NOT NULL'),
    ('ngayCheckOut', 'DATETIME', '', 'NOT NULL'),
    ('trangThai', 'NVARCHAR', '', 'NOT NULL')
])

add_db_table(doc, 6, 'Chi Tiết Đặt Phòng (ChiTietDatPhong)', [
    ('maCTDP', 'VARCHAR', 'PK', 'NOT NULL'),
    ('maPhong', 'VARCHAR', 'FK', 'REFERENCES Phong(maPhong)'),
    ('maDat', 'VARCHAR', 'FK', 'REFERENCES DatPhong(maDat)'),
    ('giaCoc', 'DECIMAL', '', ''),
    ('soNguoi', 'INT', '', '> 0'),
    ('ghiChu', 'NVARCHAR', '', '')
])

add_db_table(doc, 7, 'Hóa Đơn (HoaDon)', [
    ('maHD', 'VARCHAR', 'PK', 'NOT NULL'),
    ('maDat', 'VARCHAR', 'FK', 'REFERENCES DatPhong(maDat)'),
    ('maNV', 'VARCHAR', 'FK', 'REFERENCES NV(maNV)'),
    ('ngayTaoHD', 'DATETIME', '', 'NOT NULL'),
    ('tienPhong', 'DECIMAL', '', ''),
    ('tienDV', 'DECIMAL', '', ''),
    ('tienCoc', 'DECIMAL', '', ''),
    ('thueVAT', 'DECIMAL', '', ''),
    ('tongTien', 'DECIMAL', '', ''),
    ('doanhThu', 'DECIMAL', '', ''),
    ('loaiHD', 'NVARCHAR', '', ''),
    ('trangThaiThanhToan', 'NVARCHAR', '', ''),
    ('phuongThucThanhToan', 'NVARCHAR', '', ''),
    ('ngayThanhToan', 'DATETIME', '', ''),
    ('ghiChuThanhToan', 'NVARCHAR', '', '')
])

add_db_table(doc, 8, 'Chi Tiết Hóa Đơn (ChiTietHoaDon)', [
    ('maCTHD', 'VARCHAR', 'PK', 'NOT NULL'),
    ('maHD', 'VARCHAR', 'FK', 'REFERENCES HoaDon(maHD)'),
    ('maCTDP', 'VARCHAR', 'FK', 'REFERENCES ChiTietDatPhong(maCTDP)'),
    ('thoiGianLuuTru', 'DECIMAL', '', ''),
    ('thanhTien', 'DECIMAL', '', '')
])

add_db_table(doc, 9, 'Dịch vụ (DV)', [
    ('maDV', 'VARCHAR', 'PK', 'NOT NULL'),
    ('tenDV', 'NVARCHAR', '', 'NOT NULL'),
    ('gia', 'DECIMAL', '', ''),
    ('loaiDV', 'NVARCHAR', '', ''),
    ('mieuTa', 'NVARCHAR', '', ''),
    ('donVi', 'NVARCHAR', '', ''),
    ('trangThai', 'BIT', '', '')
])

add_db_table(doc, 10, 'Bảng Giá Dịch Vụ - Header (BangGiaDV_Header)', [
    ('maBangGia', 'VARCHAR', 'PK', 'NOT NULL'),
    ('tenBangGia', 'NVARCHAR', '', ''),
    ('ngayApDung', 'DATE', '', ''),
    ('ngayHetHieuLuc', 'DATE', '', 'ngayHetHieuLuc > ngayApDung'),
    ('trangThai', 'BIT', '', '')
])

add_db_table(doc, 11, 'Bảng Giá Dịch Vụ - Detail (BangGiaDV_Detail)', [
    ('maBangGia', 'VARCHAR', 'PK, FK', 'REFERENCES BangGiaDV_Header(maBangGia)'),
    ('maDV', 'VARCHAR', 'PK, FK', 'REFERENCES DV(maDV)'),
    ('giaDV', 'DECIMAL', '', '')
])

add_db_table(doc, 12, 'Dịch vụ Sử Dụng (DichVuSuDung)', [
    ('maDV', 'VARCHAR', 'PK, FK', 'REFERENCES DV(maDV)'),
    ('maCTDP', 'VARCHAR', 'PK, FK', 'REFERENCES ChiTietDatPhong(maCTDP)'),
    ('ngaySuDung', 'DATE', '', ''),
    ('soLuong', 'INT', '', '> 0'),
    ('giaDV', 'DECIMAL', '', ''),
    ('trangThai', 'BIT', '', '')
])

doc.save('DacTa_RangBuoc.docx')
print("Document generated successfully.")
